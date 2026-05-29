from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_ROOT / "docs" / "guia_presentacion_notebooks_cleaning_eda.docx"


COLORS = {
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "light_blue": "E8EEF5",
    "light_gray": "F2F4F7",
    "border": "B7C7D9",
    "muted": "595959",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for margin_name, margin_value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=COLORS["border"], size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_font(run, bold=False, italic=False, color=None, size=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def add_paragraph(document, text="", style=None, bold_prefix=None):
    paragraph = document.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        prefix_run = paragraph.add_run(bold_prefix)
        set_font(prefix_run, bold=True)
        body_run = paragraph.add_run(text[len(bold_prefix):])
        set_font(body_run)
    else:
        run = paragraph.add_run(text)
        set_font(run)
    return paragraph


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_font(run)


def add_numbered(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        run = paragraph.add_run(item)
        set_font(run)


def add_callout(document, title, body):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    set_table_borders(table, color="D9E2EC", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    paragraph = cell.paragraphs[0]
    title_run = paragraph.add_run(title)
    set_font(title_run, bold=True, color=COLORS["dark_blue"])
    paragraph.add_run("\n")
    body_run = paragraph.add_run(body)
    set_font(body_run)
    document.add_paragraph()


def add_label_detail_table(document, rows, col_widths=(2700, 6660), header=None):
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)

    if header:
        row = table.add_row()
        row.cells[0].merge(row.cells[1])
        cell = row.cells[0]
        set_cell_shading(cell, COLORS["light_blue"])
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(header)
        set_font(run, bold=True, color=COLORS["dark_blue"])

    for label, detail in rows:
        row = table.add_row()
        for idx, width in enumerate(col_widths):
            set_cell_width(row.cells[idx], width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(row.cells[0], COLORS["light_gray"])
        label_run = row.cells[0].paragraphs[0].add_run(label)
        set_font(label_run, bold=True, color=COLORS["dark_blue"])
        detail_run = row.cells[1].paragraphs[0].add_run(detail)
        set_font(detail_run)

    document.add_paragraph()
    return table


def add_matrix_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_width(cell, widths[index])
        set_cell_shading(cell, COLORS["light_blue"])
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(header)
        set_font(run, bold=True, color=COLORS["dark_blue"])

    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            run = cell.paragraphs[0].add_run(value)
            set_font(run)

    document.add_paragraph()
    return table


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 18, 10),
        ("Heading 2", 13, COLORS["blue"], 14, 7),
        ("Heading 3", 12, COLORS["dark_blue"], 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style_name in ["List Bullet", "List Number"]:
        style = styles[list_style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Guia notebooks 02 Cleaning y 03 EDA")
    set_font(run, color=COLORS["muted"], size=9)


def accent_document(document):
    replacements = {
        "Guia": "Guía",
        "exposicion": "exposición",
        "explicacion": "explicación",
        "Vision": "Visión",
        "despues": "después",
        "analisis": "análisis",
        "codigo": "código",
        "ambiguedades": "ambigüedades",
        "Que": "Qué",
        "Por que": "Por qué",
        "normalizacion": "normalización",
        "unificacion": "unificación",
        "mapeos": "mapeos",
        "duplicados": "duplicados",
        "tecnologicas": "tecnológicas",
        "validaciones": "validaciones",
        "critica": "crítica",
        "titulo": "título",
        "utiles": "útiles",
        "Utiles": "Útiles",
        "Librerias": "Librerías",
        "librerias": "librerías",
        "geografia": "geografía",
        "hibrido": "híbrido",
        "hibridas": "híbridas",
        "presencial": "presencial",
        "tecnologias": "tecnologías",
        "tecnologia": "tecnología",
        "hipotesis": "hipótesis",
        "clasificacion": "clasificación",
        "informacion": "información",
        "numérica": "numérica",
        "numero": "número",
        "comun": "común",
        "espaniol": "español",
        "espanol": "español",
        "senal": "señal",
        "Practica": "Práctica",
        "metodologia": "metodología",
        "Metodologia": "Metodología",
        "Como explicar": "Cómo explicar",
        "presentacion": "presentación",
        "Presentacion": "Presentación",
        "Limitaciones": "Limitaciones",
        "aproximacion": "aproximación",
        "discusion": "discusión",
        "estadistica": "estadística",
        "reproducible": "reproducible",
        "accionables": "accionables",
        "visualizaciones": "visualizaciones",
        "visualizaciónes": "visualizaciones",
        "comparacion": "comparación",
        "comparables": "comparables",
        "segun": "según",
        "Tambien": "También",
        "tambien": "también",
        "deberia": "debería",
        "deberian": "deberían",
        "rapida": "rápida",
        "fisicas": "físicas",
        "analitica": "analítica",
        "analiticas": "analíticas",
        "Practico": "Práctico",
        "parte mas": "parte más",
        "fuente con mas": "fuente con más",
        "ofertas con mas": "ofertas con más",
        "empresas con mas": "empresas con más",
        "puestos mas": "puestos más",
        "Industrias mas": "Industrias más",
        "Seniority mas": "Seniority más",
        "Puestos mas": "Puestos más",
        "Skills mas": "Skills más",
        "Tecnologias mas": "Tecnologías más",
        "código sea mas": "código sea más",
        "rol mas": "rol más",
        "ubicacion limpia mas": "ubicación limpia más",
        "modalidad mas": "modalidad más",
        "skill mas": "skill más",
        "tecnología mas": "tecnología más",
        "graficos estadisticos mas": "gráficos estadísticos más",
        "con mas o menos": "con más o menos",
    }

    def replace_in_paragraph(paragraph):
        for run in paragraph.runs:
            text = run.text
            for old, new in replacements.items():
                text = text.replace(old, new)
            run.text = text

    for paragraph in document.paragraphs:
        replace_in_paragraph(paragraph)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)


def build_document():
    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Guia para presentar los notebooks 02 Cleaning y 03 EDA")
    set_font(run, bold=True, color=COLORS["dark_blue"], size=24)

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run(
        "Proyecto EDA roles y datos | Documento de apoyo para exposicion"
    )
    set_font(subtitle_run, italic=True, color=COLORS["muted"], size=11)

    add_callout(
        document,
        "Idea principal para abrir la explicacion",
        "El cleaning deja todas las fuentes preparadas y comparables; el EDA usa esos datos ya unificados para entender la estructura, la calidad y los primeros patrones del mercado analizado.",
    )

    document.add_heading("1. Vision general del flujo", level=1)
    add_paragraph(
        document,
        "Los dos notebooks forman una secuencia: primero se limpian y estandarizan los datos; despues se exploran para detectar patrones y preparar las visualizaciones.",
    )
    add_numbered(
        document,
        [
            "02_cleaning.ipynb transforma datasets crudos en archivos limpios y consistentes.",
            "03_eda.ipynb analiza esos archivos limpios para entender roles, fuentes, nulos, salarios, skills y tecnologias.",
            "04_visualizations.ipynb deberia apoyarse en las conclusiones y columnas generadas por estos dos notebooks.",
        ],
    )

    add_label_detail_table(
        document,
        [
            ("Notebook 02", "Limpieza, normalizacion, unificacion de columnas, datasets derivados y validaciones finales."),
            ("Notebook 03", "Analisis exploratorio: estructura, calidad, distribuciones principales, salarios, skills, tecnologias y conclusiones iniciales."),
            ("Dataset principal", "jobs_all_clean.csv, porque unifica las ofertas de varias fuentes bajo nombres de columnas comunes."),
            ("Regla de consistencia", "Columnas en ingles y snake_case; markdown en espanol; comentarios bilingues; variables y funciones en ingles."),
        ],
        header="Resumen de la arquitectura de trabajo",
    )

    document.add_heading("2. Notebook 02 - Cleaning", level=1)
    add_paragraph(
        document,
        "Este notebook tiene como objetivo convertir los datos crudos en datasets limpios, estandarizados y listos para analisis. La parte mas importante es que unifica nombres de columnas para que todas las fuentes puedan compararse sin ambiguedades.",
    )

    add_matrix_table(
        document,
        ["Bloque", "Que hace", "Por que importa"],
        [
            ("Rutas e imports", "Define PROJECT_ROOT, DATA_RAW y DATA_CLEAN; carga pandas, numpy, os, re y ast.", "Garantiza que el notebook funcione desde la carpeta del proyecto o desde notebooks."),
            ("Funciones generales", "Normaliza columnas, aplica mapeos, crea columnas faltantes y limpia texto.", "Evita repetir codigo y asegura criterios de limpieza iguales."),
            ("Plantilla scraping", "Crea una estructura comun para futuras ofertas scrapeadas.", "Permite incorporar nuevas fuentes sin romper el modelo de datos."),
            ("Carga de datasets", "Lee ofertas de empleo, TecnoEmpleo y Stack Overflow.", "Separa las fuentes originales antes de transformarlas."),
            ("Limpieza por fuente", "Limpia duplicados, textos y nombres de columnas.", "Cada dataset queda usable por separado."),
            ("Unificacion de ofertas", "Construye jobs_all_clean con columnas comunes.", "Es la base principal para el EDA."),
            ("Skills y salarios", "Genera job_skills_long y salary_clean.", "Permite rankings de skills y analisis salarial aproximado."),
            ("Ubicaciones", "Crea location_clean, city_clean e is_remote.", "Facilita analizar ciudades y modalidad remota."),
            ("Validaciones", "Comprueba columnas, tipos, IDs, archivos esperados y formato snake_case.", "Da confianza antes de pasar al EDA."),
        ],
        [1700, 3950, 3710],
    )

    document.add_heading("2.1 Datasets que genera cleaning", level=2)
    add_matrix_table(
        document,
        ["Archivo", "Contenido", "Uso posterior"],
        [
            ("jobs_clean.csv", "Ofertas originales limpias con columnas normalizadas.", "Analisis de la fuente de ofertas de datos."),
            ("tecno_jobs_clean.csv", "Ofertas de TecnoEmpleo traducidas a columnas estandar en ingles.", "Comparacion con otras fuentes de empleo."),
            ("jobs_all_clean.csv", "Dataset unificado de ofertas con job_id y source_dataset.", "Dataset principal del EDA."),
            ("job_skills_long.csv", "Skills en formato largo: una fila por oferta y skill.", "Ranking y comparacion de skills demandadas."),
            ("stack_tech_columns_clean.csv", "Columnas tecnologicas de Stack Overflow en snake_case.", "Base reducida para transformar tecnologias."),
            ("technologies_clean_long_format.csv", "Tecnologias en formato largo con response_id, technology, category y type.", "Analisis de tecnologias usadas y deseadas."),
            ("technology_rankings*.csv", "Rankings de tecnologias used y wanted.", "Comparar preferencias tecnicas con skills de ofertas."),
            ("cleaning_validation_summary.csv", "Resultado de validaciones automaticas.", "Comprobar que el cleaning termino correctamente."),
        ],
        [2450, 3900, 3010],
    )

    document.add_heading("2.2 Punto clave de unificacion de columnas", level=2)
    add_paragraph(
        document,
        "La mejora critica fue traducir y unificar columnas para que todos los datasets hablen el mismo idioma tecnico. Por ejemplo, TecnoEmpleo pasa de columnas como titulo, empresa, salario, ubicacion o enlace a job_title, company, salary, location y link.",
    )
    add_paragraph(
        document,
        "En Stack Overflow tambien se cambia el formato original tipo ResponseId o LanguageHaveWorkedWith a response_id y language_have_worked_with. Esto reduce errores en el EDA y hace que el codigo sea mas legible.",
    )

    document.add_heading("3. Notebook 03 - EDA", level=1)
    add_paragraph(
        document,
        "El EDA usa los archivos limpios para entender que contienen los datos y que calidad tienen. No busca demostrar una hipotesis final, sino orientar las visualizaciones y detectar patrones iniciales.",
    )

    add_matrix_table(
        document,
        ["Seccion", "Analisis realizado", "Mensaje para explicar"],
        [
            ("Carga y validacion", "Carga todos los CSV limpios y revisa cleaning_validation_summary.", "No analizamos datos crudos; partimos de datos ya validados."),
            ("Estructura", "Resume filas, columnas y nombres de variables.", "Sirve para saber que aporta cada archivo."),
            ("Nulos", "Calcula missing_values y missing_pct por columna.", "Indica que variables son fiables y cuales requieren cautela."),
            ("Fuentes", "Cuenta ofertas por source_dataset.", "Evita que una fuente dominante sesgue la interpretacion."),
            ("Puestos", "Ranking de job_title y clasificacion por job_family.", "Agrupa roles para leer tendencias generales."),
            ("Empresas", "Top empresas y concentracion de ofertas.", "Comprueba si pocas empresas dominan la muestra."),
            ("Ubicacion y modalidad", "Analiza city_clean, is_remote y work_modality.", "Permite hablar de geografia y remoto/hibrido/presencial."),
            ("Seniority e industria", "Frecuencias de seniority_level e industry.", "Son variables parciales porque no todas las fuentes las informan."),
            ("Salarios", "Disponibilidad y distribucion de salary_clean.", "salary_clean es aproximado, util para exploracion pero no exacto."),
            ("Skills", "Ranking de skills desde job_skills_long.", "Permite ver demanda tecnica en ofertas."),
            ("Tecnologias", "Ranking used y wanted desde Stack Overflow.", "Complementa las ofertas con datos de comunidad profesional."),
            ("Fechas", "Parseo de post_date y ofertas por mes/fuente.", "Permite revisar actividad temporal cuando la fecha es valida."),
        ],
        [1850, 3800, 3710],
    )

    document.add_heading("3.1 Resultados iniciales del EDA", level=2)
    add_bullets(
        document,
        [
            "El dataset unificado original contiene 1542 ofertas y 17 columnas.",
            "La fuente con mas ofertas es df_jobs, con 942 registros.",
            "La familia de rol mas frecuente sale como data_science_ai.",
            "La ubicacion limpia mas frecuente es Madrid.",
            "La modalidad mas frecuente aparece como unknown porque muchas ofertas no indican claramente remoto, hibrido o presencial.",
            "salary_clean esta disponible en el 69.46% de las ofertas.",
            "La skill mas frecuente en ofertas es python.",
            "La tecnologia mas deseada en Stack Overflow es openAI GPT (chatbot models), dentro de ai_model_tool.",
        ],
    )

    document.add_heading("4. Librerias utilizadas y para que sirven", level=1)
    add_matrix_table(
        document,
        ["Libreria", "Donde aparece", "Para que sirve en el proyecto"],
        [
            ("pandas", "Cleaning y EDA", "Cargar CSV, transformar DataFrames, calcular nulos, rankings, agrupaciones y exportar archivos limpios."),
            ("numpy", "Cleaning y EDA", "Representar valores nulos con np.nan y hacer calculos numericos como medias salariales."),
            ("pathlib", "Cleaning y EDA", "Construir rutas robustas a data/raw y data/clean sin depender de strings fragiles."),
            ("os", "Cleaning", "Comprobaciones sencillas de archivos y directorios."),
            ("re", "Cleaning y EDA", "Usar expresiones regulares para normalizar columnas, limpiar salarios, detectar ubicaciones y clasificar texto."),
            ("ast", "Cleaning", "Interpretar listas de skills guardadas como texto cuando vienen en formato tipo Python list."),
            ("matplotlib.pyplot", "EDA", "Crear figuras base para histogramas, boxplots, barras y lineas."),
            ("seaborn", "EDA", "Generar graficos estadisticos mas legibles con menos codigo."),
            ("IPython.display", "EDA", "Mostrar tablas de forma clara dentro del notebook."),
        ],
        [1900, 1850, 5610],
    )

    document.add_heading("5. Como explicar la metodologia", level=1)
    add_numbered(
        document,
        [
            "Primero se revisaron las fuentes crudas y se identifico que no tenian los mismos nombres de columnas.",
            "Se creo una capa de limpieza reutilizable: normalizar nombres, mapear aliases, limpiar texto y crear columnas faltantes.",
            "Se estandarizaron todas las columnas principales en ingles y snake_case para facilitar continuidad entre notebooks.",
            "Se construyo un dataset unificado de ofertas, manteniendo source_dataset para no perder el origen de cada registro.",
            "Se generaron variables derivadas para el analisis: salary_clean, city_clean, is_remote, job_skills_long y rankings tecnologicos.",
            "En el EDA se partio de los archivos limpios, no de los crudos, y se analizaron estructura, calidad y patrones iniciales.",
            "Las conclusiones se trataron como exploratorias, porque algunas variables tienen nulos o proceden de fuentes con alcances distintos.",
        ],
    )

    document.add_heading("6. Frases utiles para la presentacion", level=1)
    add_label_detail_table(
        document,
        [
            ("Sobre cleaning", "La limpieza no solo elimina duplicados; tambien crea una estructura comun para que las fuentes sean comparables."),
            ("Sobre columnas", "La decision de usar nombres en ingles y snake_case reduce errores y mantiene continuidad en todo el proyecto."),
            ("Sobre jobs_all_clean", "Es el dataset principal porque concentra las ofertas de empleo en una misma estructura y conserva la fuente original."),
            ("Sobre salary_clean", "No es un salario exacto, sino una aproximacion numerica para poder explorar rangos y distribuciones."),
            ("Sobre nulos", "Los nulos no son solo un problema: tambien nos dicen que variables se pueden analizar con mas o menos confianza."),
            ("Sobre modalidad", "El alto peso de unknown indica que muchas ofertas no comunican claramente si son remotas, hibridas o presenciales."),
            ("Sobre skills", "job_skills_long permite pasar de una celda con muchas skills a una estructura analizable, una skill por fila."),
            ("Sobre Stack Overflow", "No es una fuente de ofertas, sino una referencia complementaria sobre tecnologias usadas y deseadas por profesionales."),
        ],
        header="Mensajes listos para defender",
    )

    document.add_heading("7. Limitaciones que conviene mencionar", level=1)
    add_bullets(
        document,
        [
            "Las fuentes tienen naturalezas diferentes: ofertas de empleo frente a respuestas de Stack Overflow.",
            "Algunas variables no aparecen en todas las fuentes, especialmente skills, seniority, industria y modalidad.",
            "La limpieza salarial es aproximada y depende del texto disponible en cada oferta.",
            "La clasificacion de familias de rol y modalidad se basa en reglas simples de texto, no en un modelo avanzado.",
            "El EDA describe patrones iniciales y sirve para guiar visualizaciones, no para hacer inferencia estadistica definitiva.",
        ],
    )

    document.add_heading("8. Cierre recomendado", level=1)
    add_callout(
        document,
        "Cierre para decir en voz alta",
        "Con estos dos notebooks dejamos una base reproducible: el cleaning convierte fuentes distintas en datos comparables, y el EDA convierte esos datos limpios en primeras conclusiones accionables para decidir que visualizar y que interpretar en la fase final.",
    )

    accent_document(document)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    output = build_document()
    print(output)
